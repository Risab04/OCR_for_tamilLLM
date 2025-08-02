# web scapping code 
import requests
from bs4 import BeautifulSoup
from urllib.parse import urljoin, urlparse
import os
import time
from docx import Document
from docx.shared import Inches
import re
from collections import deque
import logging
 
class WebScraper:
    def __init__(self, base_url, output_folder="scraped_content"):
        self.base_url = base_url
        self.output_folder = output_folder
        self.visited_urls = set()
        self.session = requests.Session()
        self.session.headers.update({
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        })
       
        # Create output folder if it doesn't exist
        os.makedirs(output_folder, exist_ok=True)
       
        # Setup logging
        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
        self.logger = logging.getLogger(__name__)
   
    def is_valid_url(self, url):
        """Check if URL is valid and belongs to the same domain"""
        try:
            parsed = urlparse(url)
            base_parsed = urlparse(self.base_url)
            return (parsed.netloc == base_parsed.netloc and
                    parsed.scheme in ['http', 'https'] and
                    url not in self.visited_urls)
        except:
            return False
   
    def clean_filename(self, filename):
        """Clean filename for Windows/Linux compatibility"""
        # Remove invalid characters
        filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
        # Limit length
        if len(filename) > 200:
            filename = filename[:200]
        return filename.strip()
   
    def extract_text_content(self, soup):
        """Extract meaningful text content from BeautifulSoup object"""
        # Remove script and style elements
        for script in soup(["script", "style", "nav", "header", "footer", "aside"]):
            script.decompose()
       
        # Get text content
        text_content = []
       
        # Extract title
        title = soup.find('title')
        if title:
            text_content.append(f"TITLE: {title.get_text().strip()}")
       
        # Extract headings and paragraphs
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'div', 'article', 'section']):
            text = element.get_text().strip()
            if text and len(text) > 20:  # Only include substantial text
                text_content.append(text)
       
        return text_content
   
    def scrape_page(self, url):
        """Scrape a single page and return content and links"""
        try:
            self.logger.info(f"Scraping: {url}")
            response = self.session.get(url, timeout=10)
            response.raise_for_status()
           
            soup = BeautifulSoup(response.content, 'html.parser')
           
            # Extract text content
            text_content = self.extract_text_content(soup)
           
            # Extract all links for further crawling
            links = []
            for link in soup.find_all('a', href=True):
                full_url = urljoin(url, link['href'])
                if self.is_valid_url(full_url):
                    links.append(full_url)
           
            return {
                'url': url,
                'title': soup.find('title').get_text().strip() if soup.find('title') else 'No Title',
                'content': text_content,
                'links': links
            }
       
        except Exception as e:
            self.logger.error(f"Error scraping {url}: {str(e)}")
            return None
   
    def create_word_document(self, page_data):
        """Create a Word document from page data"""
        try:
            doc = Document()
           
            # Add title
            title = doc.add_heading(page_data['title'], 0)
           
            # Add URL
            doc.add_paragraph(f"URL: {page_data['url']}")
            doc.add_paragraph("=" * 50)
           
            # Add content
            for content in page_data['content']:
                # Check if it's a heading (starts with common heading patterns)
                if any(content.upper().startswith(prefix) for prefix in ['TITLE:', 'CHAPTER', 'SECTION']):
                    doc.add_heading(content, level=1)
                else:
                    doc.add_paragraph(content)
           
            # Save document
            filename = self.clean_filename(f"{page_data['title']}.docx")
            filepath = os.path.join(self.output_folder, filename)
           
            # Handle duplicate filenames
            counter = 1
            original_filepath = filepath
            while os.path.exists(filepath):
                name, ext = os.path.splitext(original_filepath)
                filepath = f"{name}_{counter}{ext}"
                counter += 1
           
            doc.save(filepath)
            self.logger.info(f"Saved: {filepath}")
            return filepath
       
        except Exception as e:
            self.logger.error(f"Error creating document for {page_data['url']}: {str(e)}")
            return None
   
    def scrape_website(self, max_pages=50, delay=1):
        """Scrape the entire website"""
        urls_to_visit = deque([self.base_url])
        scraped_pages = []
       
        while urls_to_visit and len(scraped_pages) < max_pages:
            current_url = urls_to_visit.popleft()
           
            if current_url in self.visited_urls:
                continue
           
            self.visited_urls.add(current_url)
           
            # Scrape the page
            page_data = self.scrape_page(current_url)
           
            if page_data:
                scraped_pages.append(page_data)
               
                # Create Word document
                self.create_word_document(page_data)
               
                # Add new links to queue
                for link in page_data['links']:
                    if link not in self.visited_urls:
                        urls_to_visit.append(link)
           
            # Be respectful to the server
            time.sleep(delay)
       
        return scraped_pages
   
    def create_master_document(self, scraped_pages):
        """Create a master document with all content"""
        try:
            master_doc = Document()
            master_doc.add_heading('Website Content Master Document', 0)
            master_doc.add_paragraph(f"Base URL: {self.base_url}")
            master_doc.add_paragraph(f"Total Pages: {len(scraped_pages)}")
            master_doc.add_paragraph("=" * 80)
           
            for i, page_data in enumerate(scraped_pages, 1):
                master_doc.add_page_break()
                master_doc.add_heading(f"Page {i}: {page_data['title']}", 1)
                master_doc.add_paragraph(f"URL: {page_data['url']}")
                master_doc.add_paragraph("-" * 50)
               
                for content in page_data['content']:
                    master_doc.add_paragraph(content)
           
            master_filepath = os.path.join(self.output_folder, "master_document.docx")
            master_doc.save(master_filepath)
            self.logger.info(f"Master document saved: {master_filepath}")
           
        except Exception as e:
            self.logger.error(f"Error creating master document: {str(e)}")
 
def main():
    # Configuration - SET YOUR URL HERE
    website_url = "https://www.bbc.com/tamil"  # Replace with your actual URL
    output_folder = r"C:\Users\jainr\OneDrive\Desktop\1_internship\TAMIL OCR\output"  # Change folder name if needed
    max_pages = 50  # Maximum pages to scrape
    delay = 1  # Delay between requests in seconds
   
    # Uncomment these lines if you want interactive input instead:
    # website_url = input("Enter the website URL: ").strip()
    # output_folder = input("Enter output folder name (default: scraped_content): ").strip() or "scraped_content"
    # max_pages = int(input("Enter maximum pages to scrape (default: 50): ") or "50")
    # delay = float(input("Enter delay between requests in seconds (default: 1): ") or "1")
   
    # Create scraper instance
    scraper = WebScraper(website_url, output_folder)
   
    print(f"\nStarting to scrape: {website_url}")
    print(f"Output folder: {output_folder}")
    print(f"Max pages: {max_pages}")
    print(f"Delay: {delay} seconds")
    print("-" * 50)
   
    # Start scraping
    scraped_pages = scraper.scrape_website(max_pages=max_pages, delay=delay)
   
    # Create master document
    if scraped_pages:
        scraper.create_master_document(scraped_pages)
        print(f"\nScraping completed successfully!")
        print(f"Total pages scraped: {len(scraped_pages)}")
        print(f"Files saved in: {output_folder}")
    else:
        print("No pages were scraped successfully.")
 
if __name__ == "__main__":

    main()
