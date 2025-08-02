# -*- coding: utf-8 -*-
"""OCR_evaluation.ipynb


Original file is located at
    https://colab.research.google.com/drive/1ewAeAOjGev_4tvpFmOC1IUalyBPvSUnk
"""

# 🔧 Install Poppler
!apt-get install -y poppler-utils

# Tesseract with splitting the page

# ✅ STEP 1: Install Required Libraries
!apt install tesseract-ocr
!apt install tesseract-ocr-tam  # Tamil language pack
!pip install pytesseract pdf2image Pillow

# ✅ STEP 2: Import Libraries
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from google.colab import files
import cv2
import numpy as np
import io

pytesseract.pytesseract.tesseract_cmd = '/usr/bin/tesseract'

# ✅ STEP 3: Upload Your PDF File
uploaded = files.upload()
pdf_file = next(iter(uploaded))

# ✅ STEP 4: Convert PDF Pages to Images
pages = convert_from_path(pdf_file, dpi=150)

print(f"✅ Total pages extracted: {len(pages)}")

# ✅ STEP 5: Function to Split and Extract Tamil Text
def process_page(page_img, page_number):
    # Convert to OpenCV format
    img_np = np.array(page_img)
    img_cv = cv2.cvtColor(img_np, cv2.COLOR_RGB2BGR)

    h, w = img_cv.shape[:2]
    mid = w // 2

    left_img = img_cv[:, :mid]
    right_img = img_cv[:, mid:]

    # Convert to RGB for Tesseract
    left_rgb = cv2.cvtColor(left_img, cv2.COLOR_BGR2RGB)
    right_rgb = cv2.cvtColor(right_img, cv2.COLOR_BGR2RGB)

    left_text = pytesseract.image_to_string(left_rgb, lang='tam+eng')
    right_text = pytesseract.image_to_string(right_rgb, lang='tam+eng')


    print(f"\n📖 Page {page_number} - LEFT SIDE TEXT:\n{'-'*40}\n{left_text.strip()}\n")
    print(f"📖 Page {page_number} - RIGHT SIDE TEXT:\n{'-'*40}\n{right_text.strip()}\n")

# ✅ STEP 6: Process Each Page
for i, page in enumerate(pages):
    process_page(page, i + 1)

# tesseract without the split
# ✅ STEP 1: Install required packages
!apt-get update -qq
!apt-get install -y -qq poppler-utils tesseract-ocr tesseract-ocr-tam
!pip install pdf2image pytesseract -q

# ✅ STEP 2: Upload Tamil PDF
from google.colab import files
print("📤 Please upload a Tamil PDF file...")
uploaded = files.upload()
pdf_path = next(iter(uploaded))

# ✅ STEP 3: Convert PDF pages to images
from pdf2image import convert_from_path
from PIL import Image
import numpy as np

print("📄 Converting PDF to images...")
pdf_images = convert_from_path(pdf_path, dpi=150, first_page=1, last_page=10)
print(f"✅ Converted {len(pdf_images)} pages")

# ✅ STEP 4: Initialize result holders
extracted_text = []

# ✅ STEP 5: Tesseract OCR function
import pytesseract

def extract_tamil_text(img):
    """Extract Tamil text using Tesseract with multiple configurations"""
    configs = [
        '--psm 6 -l tam',      # Uniform block of text
        '--psm 3 -l tam',      # Fully automatic page segmentation
        '--psm 6 -l tam+eng',  # Tamil + English
        '--psm 8 -l tam',      # Single word
        '--psm 4 -l tam',      # Single column of text
    ]

    best_text = ""
    for config in configs:
        try:
            text = pytesseract.image_to_string(img, config=config)
            if text.strip() and len(text.strip()) > len(best_text):
                best_text = text.strip()
        except Exception as e:
            print(f"⚠️ Config '{config}' failed: {e}")
            continue

    return best_text

# ✅ STEP 6: Process only first 10 pages
pages_to_process = pdf_images[:10]  # Extract only first 10 pages
print(f"📋 Processing first {len(pages_to_process)} pages for testing...")

for page_num, img in enumerate(pages_to_process, start=1):
    print(f"\n📝 Processing Page {page_num}/{len(pdf_images)}...")

    # Convert to RGB format
    img = img.convert("RGB")

    # Extract text using Tesseract
    print("🔍 Running Tesseract OCR...")
    try:
        page_text = extract_tamil_text(img)
        if page_text:
            extracted_text.append(f"=== Page {page_num} ===\n{page_text}\n")
            print(f"✅ Extracted {len(page_text)} characters from page {page_num}")
        else:
            print(f"⚠️ No text found on page {page_num}")
    except Exception as e:
        print(f"❌ Failed to process page {page_num}: {e}")

# ✅ STEP 7: Compile and save results
if extracted_text:
    final_text = "\n".join(extracted_text)
    print(f"\n✅ OCR Complete! Extracted text from {len(extracted_text)} pages")
else:
    final_text = "⚠️ No text could be extracted. Try a higher quality scan or check if the PDF contains Tamil text."

# ✅ STEP 8: Save output file
output_file = "tamil_pdf_tesseract_output.txt"
with open(output_file, "w", encoding="utf-8") as f:
    f.write("Tamil PDF OCR Results (Tesseract)\n")
    f.write("="*50 + "\n")
    f.write(f"PDF File: {pdf_path}\n")
    f.write(f"Pages Processed: {len(pdf_images)}\n")
    f.write(f"OCR Engine: Tesseract\n\n")
    f.write("Extracted Text:\n")
    f.write("-"*30 + "\n")
    f.write(final_text)

# ✅ STEP 9: Download result
files.download(output_file)
print(f"\n📁 Download ready: {output_file}")
print("\n📋 Preview of extracted text:")
print("-" * 30)
print(final_text[:500] + "..." if len(final_text) > 500 else final_text)

# ✅ STEP 1: Install required packages
!apt-get update -qq
!apt-get install -y -qq poppler-utils tesseract-ocr tesseract-ocr-tam libgl1-mesa-glx
!pip install pdf2image pytesseract opencv-python-headless -q

# ✅ STEP 2: Upload Tamil PDF
from google.colab import files
print("📤 Please upload a Tamil PDF file...")
uploaded = files.upload()
pdf_path = next(iter(uploaded))

# ✅ STEP 3: Convert PDF pages to images
from pdf2image import convert_from_path
from PIL import Image
import numpy as np

print("📄 Converting PDF to images...")
pdf_images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=10)
print(f"✅ Converted {len(pdf_images)} pages")

# ✅ STEP 4: Initialize result holders
extracted_text = []
extracted_tables = []

# ✅ STEP 5: Tesseract OCR for Tamil + English
import pytesseract

def extract_tamil_text(img):
    configs = [
        '--psm 6 -l tam',
        '--psm 3 -l tam',
        '--psm 6 -l tam+eng',
        '--psm 8 -l tam',
        '--psm 4 -l tam',
    ]
    best_text = ""
    for config in configs:
        try:
            text = pytesseract.image_to_string(img, config=config)
            if text.strip() and len(text.strip()) > len(best_text):
                best_text = text.strip()
        except Exception as e:
            print(f"⚠️ Config '{config}' failed: {e}")
            continue
    return best_text

# ✅ STEP 6: Table Detection using OpenCV
import cv2

def detect_table_and_ocr(image_pil):
    image_cv = np.array(image_pil.convert('L'))  # Grayscale
    _, binary = cv2.threshold(image_cv, 180, 255, cv2.THRESH_BINARY_INV)

    # Horizontal lines
    horizontal_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (30, 1))
    detect_horizontal = cv2.morphologyEx(binary, cv2.MORPH_OPEN, horizontal_kernel, iterations=2)

    # Vertical lines
    vertical_kernel = cv2.getStructuringElement(cv2.MORPH_RECT, (1, 30))
    detect_vertical = cv2.morphologyEx(binary, cv2.MORPH_OPEN, vertical_kernel, iterations=2)

    # Combine masks
    table_mask = cv2.addWeighted(detect_horizontal, 0.5, detect_vertical, 0.5, 0.0)
    contours, _ = cv2.findContours(table_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)

    table_data = []
    for idx, cnt in enumerate(contours):
        x, y, w, h = cv2.boundingRect(cnt)
        if w > 100 and h > 50:
            roi = image_pil.crop((x, y, x + w, y + h))
            text = pytesseract.image_to_string(roi, config='--psm 6 -l tam+eng')
            if text.strip():
                table_data.append(f"--- Table {idx + 1} ---\n{text.strip()}")
    return table_data

# ✅ STEP 7: Process first 10 pages
pages_to_process = pdf_images[:10]
print(f"📋 Processing first {len(pages_to_process)} pages...")

for page_num, img in enumerate(pages_to_process, start=1):
    print(f"\n📝 Processing Page {page_num}/{len(pdf_images)}...")
    img = img.convert("RGB")

    # OCR for text
    print("🔍 Running OCR for body text...")
    try:
        page_text = extract_tamil_text(img)
        if page_text:
            extracted_text.append(f"=== Page {page_num} ===\n{page_text}\n")
            print(f"✅ Extracted {len(page_text)} characters")
        else:
            print(f"⚠️ No body text found on page {page_num}")
    except Exception as e:
        print(f"❌ Text OCR failed on page {page_num}: {e}")

    # OCR for tables
    print("📊 Detecting tables...")
    try:
        tables = detect_table_and_ocr(img)
        if tables:
            extracted_tables.extend([f"=== Page {page_num} ===\n{tbl}\n" for tbl in tables])
            print(f"✅ Extracted {len(tables)} tables")
        else:
            print(f"⚠️ No tables found on page {page_num}")
    except Exception as e:
        print(f"❌ Table detection failed on page {page_num}: {e}")

# ✅ STEP 8: Compile final output
final_text = "\n".join(extracted_text) if extracted_text else "⚠️ No text extracted."
final_tables = "\n".join(extracted_tables) if extracted_tables else "⚠️ No tables extracted."

output_file = "tamil_pdf_with_text_and_tables.txt"
with open(output_file, "w", encoding="utf-8") as f:
    f.write("Tamil PDF OCR Results (Text + Tables)\n")
    f.write("="*60 + "\n")
    f.write(f"PDF File: {pdf_path}\n")
    f.write(f"Pages Processed: {len(pages_to_process)}\n")
    f.write("OCR Engine: Tesseract\n\n")
    f.write("🔍 Extracted Body Text:\n" + "-"*30 + "\n")
    f.write(final_text)
    f.write("\n\n📊 Extracted Tables:\n" + "-"*30 + "\n")
    f.write(final_tables)

# ✅ STEP 9: Download result
files.download(output_file)
print(f"\n📁 Download ready: {output_file}")
print("\n📋 Preview of extracted text:")
print("-" * 30)
print(final_text[:500] + "..." if len(final_text) > 500 else final_text)

###
###

# Surya OCR
!pip install surya-ocr
!apt-get update && apt-get install -y poppler-utils

from google.colab import files
uploaded = files.upload()
pdf_path = list(uploaded.keys())[0]

!surya_ocr "{pdf_path}" --page_range 1-8 --langs ta --output_dir output --images

from surya.recognition import RecognitionPredictor
from surya.detection import DetectionPredictor
from pdf2image import convert_from_path

# Convert first 8 pages
images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=4)

# Load predictors
det = DetectionPredictor()
rec = RecognitionPredictor()

# Run OCR on each page
for i, img in enumerate(images, start=1):
    preds = rec([img], det_predictor=det)
    print(f"\n📄 Page {i} Tamil Text:")
    for line in preds[0].text_lines:
        print(line.text)
###
###
# easy OCR
!pip install easyocr pdf2image

from google.colab import files
uploaded = files.upload()
pdf_path = list(uploaded.keys())[0]

from pdf2image import convert_from_path
images = convert_from_path(pdf_path, dpi=300, first_page=1, last_page=8)

import easyocr
from PIL import Image

# Initialize Tamil OCR reader
reader = easyocr.Reader(['en', 'ta'])  # Add English to restore expected character count

# OCR and print
for i, img in enumerate(images, start=1):
    img_path = f"page_{i}.jpg"
    img.save(img_path, "JPEG")
    print(f"\n📄 Page {i} Tamil Text:")

    results = reader.readtext(img_path, detail=0, paragraph=True)
    for line in results:
        print(line)

# with open("easyocr_tamil_output.txt", "w", encoding="utf-8") as f:
#     for i, img in enumerate(images, start=1):
#         img_path = f"page_{i}.jpg"
#         img.save(img_path, "JPEG")
#         results = reader.readtext(img_path, detail=0, paragraph=True)
#         f.write(f"\n📄 Page {i} Tamil Text:\n")
#         for line in results:
#             f.write(line + "\n")

# # Tamil PDF OCR with PaddleOCR - Google Colab
# # This notebook extracts Tamil text from the first 5 pages of a PDF

# # Step 1: Install required packages
# !pip install paddlepaddle paddleocr pdf2image pillow

# # Step 2: Install poppler-utils (required for pdf2image)
# !apt-get update
# !apt-get install -y poppler-utils

# Step 3: Import libraries
import os
import cv2
import numpy as np
from paddleocr import PaddleOCR
from pdf2image import convert_from_path
from PIL import Image
import matplotlib.pyplot as plt
from google.colab import files
import tempfile

# Step 4: Initialize PaddleOCR with Tamil language support
# Initialize OCR with Tamil language
try:
    ocr = PaddleOCR(use_angle_cls=True, lang='ta', use_gpu=False)
    print("PaddleOCR initialized successfully with Tamil language support")
except Exception as e:
    print(f"Error initializing PaddleOCR: {e}")
    # Fallback initialization
    ocr = PaddleOCR(lang='ta')
    print("PaddleOCR initialized with fallback settings")

def extract_tamil_text_from_pdf(pdf_path, max_pages=5):
    """
    Extract Tamil text from first 5 pages of PDF using PaddleOCR

    Args:
        pdf_path: Path to the PDF file
        max_pages: Maximum number of pages to process (default: 5)

    Returns:
        Dictionary with page numbers and extracted text
    """

    print(f"Converting PDF to images (first {max_pages} pages)...")

    # Convert PDF to images (first 5 pages only) with higher DPI for better quality
    try:
        images = convert_from_path(pdf_path, first_page=1, last_page=max_pages, dpi=400)
        print(f"Successfully converted {len(images)} pages to images")
    except Exception as e:
        print(f"Error converting PDF: {e}")
        return None

    extracted_text = {}

    # Process each page
    for page_num, image in enumerate(images, 1):
        print(f"\nProcessing page {page_num}...")

        # Convert PIL image to numpy array
        img_array = np.array(image)

        # Preprocess image for better OCR
        # Convert to grayscale
        if len(img_array.shape) == 3:
            img_gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
        else:
            img_gray = img_array

        # Apply some image processing to improve OCR
        # Increase contrast
        img_processed = cv2.convertScaleAbs(img_gray, alpha=1.2, beta=10)

        # Apply Gaussian blur to reduce noise
        img_processed = cv2.GaussianBlur(img_processed, (1, 1), 0)

        # Perform OCR
        try:
            result = ocr.ocr(img_processed)

            # Debug: Print result structure for first page
            if page_num == 1:
                print(f"Debug - OCR result type: {type(result)}")
                print(f"Debug - OCR result length: {len(result) if result else 'None'}")
                if result and len(result) > 0:
                    print(f"Debug - First result type: {type(result[0])}")
                    if result[0] and len(result[0]) > 0:
                        print(f"Debug - First line structure: {result[0][0] if result[0] else 'None'}")

            # Extract text from OCR result
            page_text = []
            word_groups = []  # To group nearby characters into words

            if result and len(result) > 0 and result[0]:
                # Sort results by vertical position (top to bottom)
                sorted_results = sorted(result[0], key=lambda x: x[0][0][1] if x and len(x) > 0 else 0)

                current_line = []
                current_y = None
                line_threshold = 20  # Pixels threshold for same line

                for line in sorted_results:
                    if line and len(line) >= 2:
                        try:
                            # Get bounding box coordinates
                            bbox = line[0]
                            y_coord = bbox[0][1]  # Top Y coordinate

                            # Get text and confidence
                            if isinstance(line[1], (list, tuple)) and len(line[1]) >= 2:
                                text = line[1][0].strip()
                                confidence = line[1][1]
                            elif isinstance(line[1], str):
                                text = line[1].strip()
                                confidence = 1.0
                            else:
                                continue

                            if confidence > 0.3 and text:  # Lower confidence threshold
                                # Check if this text is on the same line as previous
                                if current_y is None or abs(y_coord - current_y) <= line_threshold:
                                    current_line.append(text)
                                    current_y = y_coord
                                else:
                                    # New line detected, join previous line
                                    if current_line:
                                        line_text = ' '.join(current_line)
                                        if len(line_text.strip()) > 1:  # Only add if more than 1 character
                                            page_text.append(line_text)
                                    current_line = [text]
                                    current_y = y_coord

                        except (IndexError, TypeError) as e:
                            print(f"Warning: Could not parse OCR result for one line: {e}")
                            continue

                # Don't forget the last line
                if current_line:
                    line_text = ' '.join(current_line)
                    if len(line_text.strip()) > 1:
                        page_text.append(line_text)

            # Join all text from the page
            full_page_text = '\n'.join(page_text)
            extracted_text[f'Page_{page_num}'] = full_page_text

            print(f"Page {page_num} - Extracted {len(page_text)} text lines")

        except Exception as e:
            print(f"Error processing page {page_num}: {e}")
            extracted_text[f'Page_{page_num}'] = f"Error processing page: {e}"

    return extracted_text

def display_results(extracted_text):
    """Display the extracted text in a formatted way"""
    print("\n" + "="*60)
    print("EXTRACTED TAMIL TEXT FROM PDF")
    print("="*60)

    for page_key, text in extracted_text.items():
        print(f"\n--- {page_key} ---")
        if text.strip():
            print(text)
        else:
            print("No text found on this page")
        print("-" * 40)

def save_results_to_file(extracted_text, output_filename='extracted_tamil_text.txt'):
    """Save extracted text to a file"""
    with open(output_filename, 'w', encoding='utf-8') as f:
        f.write("EXTRACTED TAMIL TEXT FROM PDF\n")
        f.write("="*60 + "\n\n")

        for page_key, text in extracted_text.items():
            f.write(f"--- {page_key} ---\n")
            f.write(text + "\n")
            f.write("-" * 40 + "\n\n")

    print(f"\nResults saved to {output_filename}")
    return output_filename

# Step 5: Upload and process PDF
print("Please upload your PDF file:")
uploaded = files.upload()

# Get the uploaded file
pdf_filename = list(uploaded.keys())[0]
print(f"\nUploaded file: {pdf_filename}")

# Step 6: Extract text from PDF
print("\nStarting Tamil text extraction...")
extracted_text = extract_tamil_text_from_pdf(pdf_filename, max_pages=5)

if extracted_text:
    # Step 7: Display results
    display_results(extracted_text)

    # Step 8: Save results to file
    output_file = save_results_to_file(extracted_text)

    # Step 9: Download the results file
    print("\nDownloading results file...")
    files.download(output_file)

    # Step 10: Display summary
    total_pages = len(extracted_text)
    total_text_length = sum(len(text) for text in extracted_text.values())

    print(f"\n" + "="*50)
    print("PROCESSING SUMMARY")
    print("="*50)
    print(f"Total pages processed: {total_pages}")
    print(f"Total characters extracted: {total_text_length}")
    print(f"Results saved to: {output_file}")

else:
    print("Failed to extract text from PDF")

# Optional: Function to process multiple PDFs
def process_multiple_pdfs():
    """Process multiple PDF files"""
    print("Upload multiple PDF files:")
    uploaded_files = files.upload()

    all_results = {}

    for filename in uploaded_files.keys():
        print(f"\nProcessing {filename}...")
        extracted_text = extract_tamil_text_from_pdf(filename, max_pages=5)
        if extracted_text:
            all_results[filename] = extracted_text

    return all_results

# Uncomment the following lines if you want to process multiple PDFs
# print("\n" + "="*60)
# print("MULTIPLE PDF PROCESSING")
# print("="*60)
# multiple_results = process_multiple_pdfs()



# # 1. Install layoutparser with layout models + pytesseract + detectron2
# !pip install -q layoutparser[layoutmodels,tesseract]
# !pip install -q "git+https://github.com/facebookresearch/detectron2.git"
# !apt install -y tesseract-ocr
# !apt install -y tesseract-ocr-tam  # Tamil language
# --- INSTALLATION SECTION ---

# Install layoutparser base
!pip install -q layoutparser

# Install layout models support (Detectron2-based)
!pip install -q "layoutparser[layoutmodels]"

# Install OCR support (includes pytesseract)
!pip install -q "layoutparser[ocr]"

# Install Tesseract and Tamil language support
!apt install -y tesseract-ocr
!apt install -y tesseract-ocr-tam

# Install Detectron2 separately (recommended for stability)
!pip install -q "git+https://github.com/facebookresearch/detectron2.git"

# 🛠️ Step 2: Import Libraries
import layoutparser as lp
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import cv2
import os
from google.colab import files
from IPython.display import display
import matplotlib.pyplot as plt
# 📂 Step 3: Upload Tamil PDF
uploaded = files.upload()
pdf_path = list(uploaded.keys())[0]
# 📸 Step 4: Convert PDF to Images
pages = convert_from_path(pdf_path, dpi=300)
os.makedirs("pdf_images", exist_ok=True)
image_paths = []
for i, page in enumerate(pages):
   img_path = f"pdf_images/page_{i+1}.png"
   page.save(img_path, "PNG")
   image_paths.append(img_path)
# 🔍 Step 5: Load LayoutParser Model
model = lp.Detectron2LayoutModel(
   config_path='lp://PubLayNet/faster_rcnn_R_50_FPN_3x/config',
   extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.5],
   label_map={0: "Text", 1: "Title", 2: "List", 3: "Table", 4: "Figure"},
)
# 🧠 Step 6: OCR Tamil text block-by-block
def extract_tamil_text_from_image(image_path):
   image = cv2.imread(image_path)
   image_rgb = cv2.cvtColor(image, cv2.COLOR_BGR2RGB)
   layout = model.detect(image_rgb)
   # Sort layout top-to-bottom
   layout = lp.Layout(sorted(layout, key=lambda b: b.block.y_1))
   md_output = []
   for block in layout:
       x_1, y_1, x_2, y_2 = map(int, block.coordinates)
       segment = image_rgb[y_1:y_2, x_1:x_2]
       pil_img = Image.fromarray(segment)
       # Tamil OCR using Tesseract (lang=tam)
       text = pytesseract.image_to_string(pil_img, lang="tam")
       label = block.type.lower()
       if label == "title":
           md_output.append(f"# {text.strip()}")
       elif label == "list":
           lines = text.strip().split("\n")
           md_output.extend([f"- {line}" for line in lines if line.strip()])
       elif label == "table":
           md_output.append(f"**[Table Detected]**\n{text.strip()}")
       else:  # Paragraph/text
           md_output.append(text.strip())
   return "\n\n".join(md_output)
# 📝 Step 7: Run OCR + Layout Detection on First Page
print("Processing first page...")
markdown_output = extract_tamil_text_from_image(image_paths[0])
# 💾 Step 8: Save and Download
with open("tamil_output.md", "w", encoding="utf-8") as f:
   f.write(markdown_output)
files.download("tamil_output.md")

# ✅ INSTALLATION
!pip install -q layoutparser[layoutmodels,ocr] python-docx
!pip install -q "git+https://github.com/facebookresearch/detectron2.git"
!apt install -y tesseract-ocr tesseract-ocr-tam

# ✅ IMPORTS
import layoutparser as lp
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
import cv2, os
from google.colab import files
from docx import Document

# ✅ UPLOAD PDF
uploaded = files.upload()
pdf_path = list(uploaded.keys())[0]

# ✅ CONVERT FIRST PAGE TO IMAGE
pages = convert_from_path(pdf_path, dpi=300)
os.makedirs("pdf_images", exist_ok=True)
img_path = "pdf_images/page_1.png"
pages[0].save(img_path)

# ✅ DETECT LAYOUT BLOCKS
import layoutparser as lp

model = lp.models.Detectron2LayoutModel(
    config_path="lp://PubLayNet/faster_rcnn_R_50_FPN_3x/config",
    label_map={0: "Text", 1: "Title", 2: "List", 3: "Table", 4: "Figure"},
    extra_config=["MODEL.ROI_HEADS.SCORE_THRESH_TEST", 0.5]
)

image = cv2.imread(img_path)
layout = model.detect(image)

# ✅ WORD-LEVEL OCR
pytesseract.pytesseract.tesseract_cmd = "/usr/bin/tesseract"
ocr_data = pytesseract.image_to_data(image, lang='tam', output_type=pytesseract.Output.DICT)

# ✅ GENERATE HTML LAYOUT
html_output = """
<!DOCTYPE html>
<html><body style='position: relative; width: 1000px; font-family: sans-serif;'>\n
"""

for i in range(len(ocr_data['text'])):
    if int(ocr_data['conf'][i]) > 50 and ocr_data['text'][i].strip():
        x, y, w, h = ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i]
        word = ocr_data['text'][i]
        html_output += f"<div style='position: absolute; top: {y}px; left: {x}px; " \
                       f"width: {w}px; height: {h}px;'>{word}</div>\n"

html_output += "</body></html>"

with open("layout_output.html", "w", encoding="utf-8") as f:
    f.write(html_output)

# ✅ GENERATE SIMPLE DOCX OUTPUT
doc = Document()
doc.add_heading('Tamil OCR Output', 0)

current_line = ""
prev_y = -1
for i in range(len(ocr_data['text'])):
    word = ocr_data['text'][i].strip()
    if not word or int(ocr_data['conf'][i]) < 50:
        continue
    y = ocr_data['top'][i]
    if prev_y == -1 or abs(y - prev_y) < 10:
        current_line += word + " "
    else:
        doc.add_paragraph(current_line.strip())
        current_line = word + " "
    prev_y = y
doc.add_paragraph(current_line.strip())
doc.save("tamil_output.docx")

# ✅ DOWNLOAD OUTPUTS
files.download("layout_output.html")
files.download("tamil_output.docx")

